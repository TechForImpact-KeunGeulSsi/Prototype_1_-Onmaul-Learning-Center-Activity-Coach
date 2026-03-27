import json
import os
from pathlib import Path
from io import BytesIO

from flask import Flask, jsonify, request, send_from_directory, send_file

try:
    from openai import OpenAI
except Exception:  # pragma: no cover
    OpenAI = None

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None


BASE_DIR = Path(__file__).resolve().parent
app = Flask(__name__, static_folder=str(BASE_DIR), static_url_path="")


def _load_env_file() -> None:
    env_path = BASE_DIR / ".env"
    if not env_path.exists():
        return
    for line in env_path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        os.environ.setdefault(key.strip(), value.strip())


def _safe_json(text: str, fallback):
    try:
        return json.loads(text)
    except Exception:
        return fallback


def _openai_client():
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key or OpenAI is None:
        return None
    return OpenAI(api_key=api_key)


def _fallback_image_data_url(title: str) -> str:
    safe_title = (title or "온마을 활동").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    svg = f"""<svg xmlns='http://www.w3.org/2000/svg' width='1024' height='768'>
<defs><linearGradient id='g' x1='0' x2='1' y1='0' y2='1'><stop offset='0%' stop-color='#dff1ff'/><stop offset='100%' stop-color='#f6fbff'/></linearGradient></defs>
<rect width='100%' height='100%' fill='url(#g)'/>
<rect x='48' y='48' width='928' height='672' rx='28' fill='white' stroke='#b9dcf8' stroke-width='3'/>
<text x='80' y='130' font-size='44' fill='#0d5f94' font-family='Arial'>온마을 활동 시각자료</text>
<text x='80' y='200' font-size='36' fill='#1f2a37' font-family='Arial'>{safe_title}</text>
<text x='80' y='270' font-size='26' fill='#3f4e61' font-family='Arial'>아이들이 마을 자원을 탐구하는 장면</text>
</svg>"""
    return "data:image/svg+xml;utf8," + svg.replace("\n", "")


@app.get("/")
def root():
    return send_from_directory(BASE_DIR, "proto_1.html")


@app.get("/health")
def health():
    status = "ok" if _openai_client() else "fallback"
    return jsonify({"status": status})


@app.post("/api/recommend")
def api_recommend():
    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "다로리 마을 탐색")
    grade = payload.get("grade", "1-2학년")
    members = payload.get("members", 12)
    duration = payload.get("duration", 40)
    place = payload.get("place", "실외")
    resources = payload.get("resources", [])
    constraints = payload.get("constraints", "")

    fallback = {
        "items": [
            {
                "title": f"{title} 관찰 미션",
                "goal": "마을 자원을 스스로 관찰하고 질문하는 힘 기르기",
                "steps": [f"도입 5분: {place} 안전규칙", f"탐색 {max(15, int(duration)-20)}분: 관찰 활동", "정리 15분: 질문 공유"],
                "cautions": "도로 접근 금지, 2인 1조 이동",
                "tags": [", ".join(resources) if resources else "자연", f"{members}명 운영"],
            },
            {
                "title": f"{title} 작은 농업 실험실",
                "goal": "농업 기반 실험으로 데이터 읽기 기초 익히기",
                "steps": ["도입 10분", "실험 20분", "정리 10분"],
                "cautions": "도구 사용 전 교사 확인",
                "tags": ["농업", grade],
            },
            {
                "title": f"{title} 공동체 인터뷰",
                "goal": "어르신 경험을 듣고 핵심 정보를 정리하기",
                "steps": ["질문카드 작성", "인터뷰/역할극", "핵심문장 정리"],
                "cautions": "경청 및 예절 지도",
                "tags": ["공동체", "표현"],
            },
        ]
    }

    client = _openai_client()
    if not client:
        return jsonify(fallback)

    prompt = f"""
너는 농촌 방과후 활동 코치야.
입력: 제목={title}, 학년={grade}, 인원={members}, 시간={duration}, 장소={place}, 자원={resources}, 제약={constraints}
요구:
- 3개 활동안을 JSON으로 반환
- 각 항목: title, goal, steps(string[]), cautions, tags(string[])
- 현실적이고 안전수칙 포함
- 교사가 최종 판단한다는 톤
출력은 JSON만.
"""
    try:
        resp = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Return JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.7,
        )
        content = resp.choices[0].message.content or "{}"
        data = _safe_json(content, fallback)
        if "items" not in data or not isinstance(data["items"], list):
            return jsonify(fallback)
        return jsonify(data)
    except Exception:
        return jsonify(fallback)


@app.post("/api/question-coach")
def api_question_coach():
    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "마을 탐색 활동")
    grade = payload.get("grade", "1-2학년")
    fallback = {
        "questions": [
            f"{title}에서 가장 신기했던 점은?",
            f"{grade} 친구에게 쉽게 설명하면?",
            "AI 답을 확인하려면 무엇을 더 찾아봐야 할까?",
            "오늘 활동이 마을 문제와 어떻게 연결될까?",
            "다음 시간에 더 탐구할 질문은?",
        ]
    }
    client = _openai_client()
    if not client:
        return jsonify(fallback)

    prompt = f"""
제목: {title}, 학년: {grade}
아이 탐구형 질문 5개를 JSON으로 생성해줘.
키는 questions(string[])만 사용.
"""
    try:
        resp = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Return JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.8,
        )
        content = resp.choices[0].message.content or "{}"
        data = _safe_json(content, fallback)
        if "questions" not in data or not isinstance(data["questions"], list):
            return jsonify(fallback)
        return jsonify(data)
    except Exception:
        return jsonify(fallback)


@app.post("/api/summary")
def api_summary():
    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "마을 탐색 활동")
    participation = payload.get("participation", "중")
    mood = payload.get("mood", "밝음")
    notes = payload.get("notes", "")

    fallback = {
        "summary": f"오늘은 '{title}' 활동을 진행했습니다. 참여도는 {participation}, 분위기는 {mood}이었고 {notes or '아이들이 협력적으로 참여했습니다.'}",
        "next": "다음 시간에는 오늘 나온 질문 중 1개를 선택해 소규모 탐구 활동으로 이어가겠습니다.",
    }
    client = _openai_client()
    if not client:
        return jsonify(fallback)

    prompt = f"""
활동명={title}, 참여도={participation}, 분위기={mood}, 메모={notes}
보호자 공유용 요약문 2~3문장과 다음 예고 1문장을 JSON으로 생성.
키: summary, next
"""
    try:
        resp = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Return JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.6,
        )
        content = resp.choices[0].message.content or "{}"
        data = _safe_json(content, fallback)
        if not isinstance(data.get("summary"), str) or not isinstance(data.get("next"), str):
            return jsonify(fallback)
        return jsonify(data)
    except Exception:
        return jsonify(fallback)


@app.post("/api/visuals")
def api_visuals():
    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "마을 탐색 활동")
    grade = payload.get("grade", "1-2학년")
    summary = payload.get("summary", "")
    next_plan = payload.get("next", "")

    fallback = {
        "table": [
            {"항목": "활동명", "내용": title},
            {"항목": "대상", "내용": grade},
            {"항목": "오늘 요약", "내용": summary or "학생들이 협력적으로 활동에 참여함"},
            {"항목": "다음 예고", "내용": next_plan or "질문 기반 탐구 활동"},
        ],
        "image_url": _fallback_image_data_url(title),
        "image_prompt": f"{title} 활동을 표현하는 교육용 일러스트",
    }

    client = _openai_client()
    if not client:
        return jsonify(fallback)

    table_prompt = f"""
활동명={title}, 대상={grade}, 요약={summary}, 다음예고={next_plan}
학부모 공유용 표를 JSON 배열로 반환해줘.
형식: {{"table":[{{"항목":"", "내용":""}}]}}
"""
    result = dict(fallback)
    try:
        table_resp = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-4o-mini"),
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "Return JSON only."},
                {"role": "user", "content": table_prompt},
            ],
            temperature=0.4,
        )
        table_content = table_resp.choices[0].message.content or "{}"
        table_data = _safe_json(table_content, {})
        if isinstance(table_data.get("table"), list):
            result["table"] = table_data["table"]
    except Exception:
        pass

    image_prompt = f"blue and white clean educational app style illustration, children rural village afterschool learning, {title}, warm friendly, no text"
    result["image_prompt"] = image_prompt
    try:
        image_resp = client.images.generate(
            model=os.environ.get("OPENAI_IMAGE_MODEL", "gpt-image-1"),
            prompt=image_prompt,
            size="1024x1024",
        )
        if image_resp.data and getattr(image_resp.data[0], "b64_json", None):
            result["image_url"] = "data:image/png;base64," + image_resp.data[0].b64_json
        elif image_resp.data and getattr(image_resp.data[0], "url", None):
            result["image_url"] = image_resp.data[0].url
    except Exception:
        pass

    return jsonify(result)


@app.post("/api/export-docx")
def api_export_docx():
    payload = request.get_json(silent=True) or {}
    title = payload.get("title", "온마을 활동")
    summary = payload.get("summary", "")
    next_plan = payload.get("next", "")
    participation = payload.get("participation", "")
    mood = payload.get("mood", "")
    notes = payload.get("notes", "")

    if Document is None:
        return jsonify({"error": "python-docx 미설치"}), 500

    doc = Document()
    doc.add_heading("온마을 활동 리포트", level=1)
    doc.add_paragraph(f"활동명: {title}")
    doc.add_paragraph(f"참여도: {participation}")
    doc.add_paragraph(f"분위기: {mood}")
    doc.add_paragraph("")
    doc.add_paragraph("오늘 요약")
    doc.add_paragraph(summary or "-")
    doc.add_paragraph("")
    doc.add_paragraph("관찰 메모")
    doc.add_paragraph(notes or "-")
    doc.add_paragraph("")
    doc.add_paragraph("다음 활동 예고")
    doc.add_paragraph(next_plan or "-")

    mem = BytesIO()
    doc.save(mem)
    mem.seek(0)
    return send_file(
        mem,
        as_attachment=True,
        download_name="onmaeul_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    _load_env_file()
    app.run(host="127.0.0.1", port=5000, debug=True)
